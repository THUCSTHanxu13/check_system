#coding:utf-8

from docx import Document
from docx.shared import Inches

class docx_utils(object):

	def __init__(self):
		pass

	def new_document(self):
		document = Document()
		return document

	def save_document(self, document, name):
		document.save(name)

	def add_heading(self, document, text, level = 1):
		document.add_heading(text, level = level)

	def add_paragraph(self, document, text, style = None):
		if style:
			para = document.add_paragraph(text, style = style)
		else:
			para = document.add_paragraph(text)
		return para

	def add_text_into_paragraph(self, paragraph, text, style = None):
		if style == None:
			paragraph.add_run(text)
		elif style == 'bold':
			paragraph.add_run(text).bold = True
		elif style == "italic":
			paragraph.add_run(text).italic = True
		else:
			paragraph.add_run(text)

	def add_page_break(self, document):
		document.add_page_break()