#coding:utf-8
from .utils import website_utils
from .utils import parser_utils
import docx
from docx import Document
import time
import random
import json
from docx.oxml.ns import qn

class docx_analyzer(object):

	def __init__(self, model_path, driver_path, upload_path, download_path, page_scope = 15, piece_scope = 30, para_scope = 25, rate_scope = 0.7):
		self.website_utils = website_utils(driver_path)
		self.parser_utils = parser_utils(model_path)
		self.page_scope = page_scope
		self.piece_scope = piece_scope
		self.para_scope = para_scope
		self.rate_scope = rate_scope
		self.upload_path = upload_path
		self.download_path = download_path

	def check(self, sen):
		res = self.website_utils.get_sogou_ab_from_text_with_scope(sen, self.page_scope)
		mx_rate = 0.0
		mx_url = ""
		mx_content = ""
		for i in res:
			rate = self.parser_utils.textcheck(sen, i[0])
			if rate > mx_rate:
				mx_rate = rate
				mx_url = i[1]
				mx_content = i[0]
		return mx_rate, mx_url, mx_content

	def piece_split(self, sen):
		pieces = sen.split("，")
		res = [pieces[0]]
		for i in pieces[1:]:
			if (len(res[-1]) + len(i) <= self.piece_scope):
				res[-1] = res[-1] + "，" + i
			else:
				res.append(i)
		return res

	def analyze_docx(self, name):
		total = 0.0
		acc = 0.0
		doc = docx.Document(name)
		all_res = []
		for p in doc.paragraphs:
			paragraph = self.parser_utils.convert_to_unicode(p.text)
			if (len(paragraph) < self.para_scope):
				all_res.append([(paragraph, "", "")])
				total += len(paragraph)
				continue
			res = []
			sentences = self.parser_utils.sentencesplit(paragraph)
			for sen in sentences:
				pieces = self.piece_split(sen)
				flag = True
				for piece in pieces:
					mx_rate, mx_url, mx_content = self.check(piece)
					total += len(piece)
					print (piece, mx_rate)
					if mx_rate > self.rate_scope:
						acc += len(piece)
						res.append((piece, mx_url, mx_content, mx_rate))
						if not flag:
							res[-1][0] = '，' + res[-1][0]
							flag = False
					else:
						res.append((piece, "", ""))
						if not flag:
							res[-1][0] = '，' + res[-1][0]
							flag = False
			all_res.append(res)
		return (all_res, acc / total)

	def export_to_html(self, res, rate):
		info = {}
		info['rate'] = "%.2f"%(rate*100)
		context = []
		fade = []
		for i in res:
			context.append([])
			for j in i:
				if j[1] != "":
					context[-1].append((True, j[0], '#'+str(len(fade))))
					fade.append(('#'+str(len(fade)),j[1],j[2], "%.2f"%(j[3]*100)))
				else:
					context[-1].append((False, j[0], ''))
		info['fade'] = fade
		info['context'] = context
		return info

	def export_to_docx(self, res, rate, name = None):
		document = Document()
		document.styles['Normal'].font.name = u'宋体'
		document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
		document.add_heading("重复率%.2f%%"%(rate*100), level = 2)
		for i in res:
			para = document.add_paragraph("    ")
			for j in i:
				if j[1] != "":
					para.add_run(j[0]).bold = True
				else:
					para.add_run(j[0])
		if name == None:
			name = time.strftime('%Y_%m_%d_%H_%M_%S',time.localtime(time.time()))
			name = name + "_%d"%((int)(random.random()*10000))
			name = self.download_path + name + '.docx'

		document.save(name)
		return name

# print (res[1])
# f = open("res.json", "w")
# f.write(json.dumps(res))
# f.close()

# f = open("res.json", "r")
# a = json.loads(f.read())
# all_res, rate = a
# f.close()