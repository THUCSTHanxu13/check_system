#coding:utf-8
from utils import website_utils
from utils import parser_utils
from docx import Document
import name
import random

class docx_analyzer(object):

	def __init__(self, model_path, driver_path, page_scope = 15, piece_scope = 30, para_scope = 25, rate_scope = 0.7):
		self.website_utils = website_utils(driver_path)
		self.parser_utils = parser_utils(model_path)
		self.page_scope = page_scope
		self.piece_scope = piece_scope
		self.para_scope = para_scope
		self.rate_scope = rate_scope

	def check(self, sen):
		res = self.website_utils.get_sogou_ab_from_text_with_scope(sen, self.page_scope)
		mx_rate = 0.0
		mx_url = ""
		mx_content = ""
		for i in res:
			rate = self.parser_utils.textcheck(sen, i[0])
			if rate > mx_rate:
				mx_rate = rate
				mx_url = i[1]
				mx_content = i[0]
		return mx_rate, mx_url, mx_content

	def piece_split(self, sen):
		pieces = sen.split("，")
		res = [pieces[0]]
		for i in pieces[1:]:
			if (len(res[-1]) + len(i) <= self.piece_scope):
				res[-1] = res[-1] + "，" + i
			else:
				res.append(i)
		return res

	def analyze_docx(self, name):
		total = 0.0
		acc = 0.0
		doc = docx.Document(name)
		all_res = []
		for p in doc.paragraphs:
			paragraph = self.parser_utils.convert_to_unicode(p.text)
			if (len(paragraph) < self.para_scope):
				all_res.append([(paragraph, "", "")])
				total += len(paragraph)
				continue
			res = []
			sentences = self.parser_utils.sentencesplit(paragraph)
			for sen in sentences:
				pieces = self.piece_split(sen)
				flag = True
				for piece in pieces:
					mx_rate, mx_url, mx_content = self.check(piece)
					total += len(piece)
					print (piece, mx_rate)
					if mx_rate > self.rate_scope:
						acc += len(piece)
						res.append((piece, mx_url, mx_content, mx_rate))
						if not flag:
							res[-1][0] = '，' + res[-1][0]
							flag = False
					else:
						res.append((piece, "", ""))
						if not flag:
							res[-1][0] = '，' + res[-1][0]
							flag = False
			all_res.append(res)
		return (all_res, acc / total)

	def export_to_html(self, res, rate):
		info = {}
		info['rate'] = rate
		context = []
		fade = []
		for i in res:
			for j in i:
				if j[1] != "":
					context.append((True, j[0], '#'+str(len(fade))))
					fade.append(('#'+str(len(fade)),j[1],j[2], "%.2f"%(rate*100)))
				else:
					context.append((False, j[0], ''))
		info['fade'] = fade
		info['context'] = context
		return info

	def export_to_docx(self, res, rate):
		new_document(self):
		document = Document()
		document.add_heading("重合率%.2f"%(rate), level = 2)
		for i in res:
			para = document.add_paragraph("    ")
			for j in i:
				if j[1] != "":
					para.add_run(j[0]).bold = True
				else:
					para.add_run(j[0])
		name = time.strftime('%Y_%m_%d_%H_%M_%S',time.localtime(time.time()))
		name = name + "_%d"%((int)(random.random()*10000))
		document.save('./download/'+name+'.docx')
		return name


analyzer = docx_analyzer(model_path = 'models', driver_path = "drivers")
res = analyzer.analyze_docx('test.docx')

print (res[1])
f = open("res.json", "w")
f.write(json.dumps(res))
f.close()



